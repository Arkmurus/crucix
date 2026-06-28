"""ARIA Document Intelligence Reade

r — 4-strategy fallback pipeline.

Cherry-picked from the v3 architecture proposal and adapted to our
existing stack (LLM provider abstraction, redis_store, config.py).

Four extraction strategies, tried in order:
  1. Standard text extraction (pdfplumber) — fast, no overhead
  2. OCR via PyMuPDF + Tesseract — handles scanned/image PDFs
  3. LLM vision model — complex layouts, handwriting (via our LLM provider)
  4. Online search for accessible version — last resort

Languages supported: English, Portuguese, French, Arabic (CPLP coverage).
Never fails silently — always reports what was attempted and why it failed.

Dependencies (gracefully degraded if missing):
  - pdfplumber (required for Strategy 1)
  - PyMuPDF / fitz (optional, Strategy 2+3)
  - pytesseract + Pillow (optional, Strategy 2)
  - python-docx (optional, .docx files)
  - trafilatura / beautifulsoup4 (optional, .html files)

Feature-gated: ARIA_DOCUMENT_READER_ENABLED env var (default ON)."""
from __future__ import annotations
from .engine_wiring import wire_success

import base64
import io
import logging
import math
import os
import asyncio
import tempfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import TYPE_CHECKING, Optional
from .wire import fail_wire  # R-F1789 §21 brain-wiring

if TYPE_CHECKING:
    from ..llm.provider import LLMProvider

logger = logging.getLogger("aria.intel.document_reader")

# ── Graceful dependency imports ─────────────────────────────────────────────

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.info("pdfplumber not installed — PDF text extraction unavailable")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = PYMUPDF_AVAILABLE  # needs both
except ImportError:
    TESSERACT_AVAILABLE = False

# R-F1104 — native .xlsx reading for DD spreadsheets
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.info("openpyxl not installed — .xlsx spreadsheet extraction unavailable")


# ── Configuration from env vars (matches existing config.py pattern) ────────

_OCR_DPI = int(os.getenv("ARIA_OCR_DPI", "300"))
_OCR_LANGUAGES = os.getenv("ARIA_OCR_LANGUAGES", "eng+por")
_VISION_MAX_PAGES = int(os.getenv("ARIA_VISION_MAX_PAGES", "10"))
_REQUEST_TIMEOUT = int(os.getenv("ARIA_DOC_TIMEOUT", "30"))

# ── Chunked vision config ─────────────────────────────────────────────────
VISION_LARGE_DOC_THRESHOLD = int(os.getenv("ARIA_VISION_LARGE_DOC_THRESHOLD", "50"))
VISION_MAX_TOTAL_PAGES = int(os.getenv("ARIA_VISION_MAX_TOTAL_PAGES", "300"))
VISION_CHUNK_SIZE = 10  # pages per chunk


# ── Feature flag ────────────────────────────────────────────────────────────

@fail_wire(module="document_reader", gap_type="file_parse")
def is_enabled() -> bool:
    val = os.getenv("ARIA_DOCUMENT_READER_ENABLED", "1") or "1"
    return val.strip().lower() not in ("0", "false", "no", "off")


# ── Data structures ─────────────────────────────────────────────────────────

@dataclass
class ExtractionResult:
    """Structured result from any document extraction attempt."""
    text: str = ""
    method: str = "UNKNOWN"
    confidence: float = 0.0
    pages_extracted: int = 0
    total_pages: int = 0
    language_detected: str = "unknown"
    gap_description: Optional[str] = None
    warnings: list[str] = field(default_factory=list)
    strategies_attempted: list[str] = field(default_factory=list)
    chunks_processed: int = 0
    chunks_total: int = 0
    chunk_failures: list = field(default_factory=list)

    @property
    def is_usable(self) -> bool:
        return self.confidence >= 0.40 and len(self.text.strip()) >= 100

    @property
    def summary(self) -> str:
        if self.is_usable:
            return (
                f"Extracted via {self.method} ({self.confidence:.0%} confidence, "
                f"{self.pages_extracted}/{self.total_pages} pages)"
            )
        if self.gap_description:
            return f"UNREADABLE: {self.gap_description}"
        return f"Extraction failed via {self.method}"


# ── Main reader ─────────────────────────────────────────────────────────────

@fail_wire(module="document_reader", gap_type="file_parse")
async def read_document(
    source: str,
    llm: "LLMProvider | None" = None,
    query: str = "",
    language_hint: str = "",
) -> ExtractionResult:
    """Read a document using the 4-strategy fallback pipeline.

    Args:
        source: File path or URL
        llm: LLM provider for vision strategy (optional)
        query: Focus query for extraction context
        language_hint: ISO language code (por, fra, ara)

    Returns: ExtractionResult with text, confidence, and metadata
    """
    if not is_enabled():
        return ExtractionResult(
            method="DISABLED",
            gap_description="Document reader disabled via ARIA_DOCUMENT_READER_ENABLED=0",
        )

    # Resolve source (download URL if needed)
    filepath = await _resolve_source(source)
    if not filepath:
        return ExtractionResult(
            method="SOURCE_RESOLUTION",
            gap_description=f"Could not resolve source: {source}",
        )

    # Detect file type
    ext = os.path.splitext(filepath)[1].lower()
    strategies_attempted: list[str] = []

    # Non-PDF fast paths
    if ext in (".txt", ".csv", ".md", ".json"):
        return _read_plaintext(filepath)
    if ext in (".html", ".htm"):
        return _read_html(filepath)
    if ext in (".docx",):
        return _read_docx(filepath)
    # R-F1104 — native .xlsx reading for DD spreadsheets
    if ext in (".xlsx", ".xls"):
        return _read_xlsx(filepath)
    if ext in (".jpg", ".jpeg", ".png", ".gif", ".webp"):
        if llm:
            return await _strategy_vision_image(filepath, llm, query)
        return ExtractionResult(
            method="IMAGE",
            gap_description="Image file requires LLM vision — no LLM configured",
        )

    # PDF pipeline — 4 strategies
    if ext != ".pdf":
        return ExtractionResult(
            method="UNSUPPORTED",
            gap_description=f"Unsupported file type: {ext}",
        )

    # Strategy 1: pdfplumber text extraction
    if PDFPLUMBER_AVAILABLE:
        strategies_attempted.append("TEXT_EXTRACTION")
        result = await asyncio.to_thread(_strategy_text_extraction, filepath)
        if result.confidence >= 0.60:
            result.strategies_attempted = strategies_attempted
            return result

    # Strategy 1b: PDF table extraction via pdfplumber
    # 2026-04-12: extracts structured tables as CSV alongside text. Critical
    # for pricing matrices, spec sheets, compliance matrices, supplier lists.
    # Only runs if Strategy 1 text extraction found tables (low text confidence
    # often means the PDF is table-heavy).
    if PDFPLUMBER_AVAILABLE and result and result.confidence < 0.80:
        strategies_attempted.append("TABLE_EXTRACTION")
        table_result = await asyncio.to_thread(_strategy_table_extraction, filepath)
        if table_result.confidence > 0 and table_result.text:
            # Merge tables into the text result
            if result and result.text:
                result.text += "\n\n" + table_result.text
                result.confidence = max(result.confidence, table_result.confidence)
                result.method += "+TABLE_EXTRACTION"
            else:
                result = table_result
            result.strategies_attempted = strategies_attempted
            if result.confidence >= 0.60:
                return result

    # Strategy 2: OCR via PyMuPDF + Tesseract
    if TESSERACT_AVAILABLE:
        strategies_attempted.append("OCR_TESSERACT")
        lang = _resolve_ocr_language(language_hint)
        result = await asyncio.to_thread(_strategy_ocr_tesseract, filepath, lang)
        if result.confidence >= 0.50:
            result.strategies_attempted = strategies_attempted
            return result
    else:
        strategies_attempted.append("OCR_UNAVAILABLE")

    # Strategy 3a: Chunked vision for large PDFs
    if llm and PYMUPDF_AVAILABLE:
        try:
            # R-F1666: page-count probe off the event loop.
            _page_count = await asyncio.to_thread(_fitz_page_count, filepath)
        except Exception:
            _page_count = 0
        if _page_count > VISION_LARGE_DOC_THRESHOLD:
            strategies_attempted.append("VISION_CHUNKED")
            result = await _read_pdf_chunked(filepath, llm, query, max_pages=VISION_MAX_TOTAL_PAGES)
            if result.confidence >= 0.40:
                result.strategies_attempted = strategies_attempted
                return result

    # Strategy 3: LLM vision model
    if llm and PYMUPDF_AVAILABLE:
        strategies_attempted.append("VISION_MODEL")
        result = await _strategy_vision_pdf(filepath, llm, query)
        if result.confidence >= 0.40:
            result.strategies_attempted = strategies_attempted
            return result

    # Strategy 4: Find accessible version online
    strategies_attempted.append("ONLINE_SEARCH")
    result = _strategy_find_online(filepath, query)
    if result.confidence >= 0.30:
        result.strategies_attempted = strategies_attempted
        return result

    # All strategies failed.
    # R-F881 (2026-05-25) — surface a TOTAL extraction failure at WARNING +
    # emit a capability_gap. Per-strategy failures above stay at debug (the
    # 4-strategy fallback is by-design — pdfplumber miss → OCR → vision →
    # online is normal), but exhausting ALL of them means the document is
    # genuinely unreadable, and pre-R-F881 that surfaced only at debug — below
    # the brain's WARNING+ error_log_handler line and invisible to the operator
    # (the silent-failure class behind the contract-upload work). Now it reaches
    # error_log_handler, the operator dashboard, and (once gap_detector is
    # reconnected) the coder.
    _r881_basename = (filepath or "?").replace("\\", "/").rsplit("/", 1)[-1]
    logger.warning(
        "[document_reader] R-F881 ALL %d strategies FAILED for %s — document "
        "UNREADABLE (tried: %s)",
        len(strategies_attempted), _r881_basename, ", ".join(strategies_attempted),
    )
    try:
        from . import capability_gaps as _cg881
        await _cg881.record_gap(
            gap_type="file_parse",
            detail=(
                f"document_reader exhausted all {len(strategies_attempted)} "
                f"strategies ({', '.join(strategies_attempted)}) on "
                f"'{_r881_basename}' — unreadable"
            ),
            source="document_reader.read_document",
        )
    except Exception as _cg_err:
        logger.debug("[document_reader] R-F881 capability_gap emit failed: %s", _cg_err)
    return ExtractionResult(
        method="ALL_STRATEGIES_FAILED",
        confidence=0.0,
        strategies_attempted=strategies_attempted,
        gap_description=(
            f"Document unreadable after {len(strategies_attempted)} strategies: "
            f"{', '.join(strategies_attempted)}. "
            f"Provide document in text-searchable format, or share the source URL."
        ),
    )


# ── Strategy 1: pdfplumber ──────────────────────────────────────────────────

def _strategy_text_extraction(filepath: str) -> ExtractionResult:
    """Extract text layer from PDF using pdfplumber."""
    try:
        pages_text = []
        total_pages = 0
        with pdfplumber.open(filepath) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)

        full_text = "\n\n".join(pages_text)
        if len(full_text.strip()) < 50:
            return ExtractionResult(
                method="TEXT_EXTRACTION", confidence=0.05,
                total_pages=total_pages, pages_extracted=len(pages_text),
                gap_description="PDF appears image-based — no text layer",
            )

        avg_chars = len(full_text) / max(total_pages, 1)
        confidence = min(0.95, avg_chars / 500)

        return ExtractionResult(
            text=full_text, method="TEXT_EXTRACTION",
            confidence=confidence, pages_extracted=len(pages_text),
            total_pages=total_pages,
        )
    except Exception as e:
        logger.debug("pdfplumber extraction failed: %s", e)
        return ExtractionResult(
            method="TEXT_EXTRACTION", confidence=0.0,
            gap_description=f"pdfplumber error: {e}",
        )


# ── Strategy 1b: PDF Table Extraction ─────────────────────────────────────
# 2026-04-12: uses pdfplumber's built-in extract_tables() to pull structured
# data from pricing matrices, spec sheets, compliance matrices, supplier lists.
# No new dependencies — pdfplumber is already required for Strategy 1.

def _strategy_table_extraction(filepath: str) -> ExtractionResult:
    """Extract tables from PDF using pdfplumber.extract_tables().

    Returns tables as CSV-formatted text blocks, one per table found.
    """
    if not PDFPLUMBER_AVAILABLE:
        return ExtractionResult(method="TABLE_EXTRACTION", confidence=0.0,
                                gap_description="pdfplumber not available")
    try:
        tables_text = []
        total_pages = 0
        with pdfplumber.open(filepath) as pdf:
            total_pages = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages[:50], 1):  # cap at 50 pages
                tables = page.extract_tables()
                if not tables:
                    continue
                for t_idx, table in enumerate(tables):
                    if not table:
                        continue
                    # Convert table rows to CSV-like format
                    rows = []
                    for row in table:
                        cells = [str(cell or "").strip().replace("\n", " ") for cell in row]
                        rows.append(" | ".join(cells))
                    if rows:
                        tables_text.append(
                            f"[TABLE p{page_num}.{t_idx + 1}]\n" + "\n".join(rows)
                        )

        if not tables_text:
            return ExtractionResult(
                method="TABLE_EXTRACTION", confidence=0.0,
                total_pages=total_pages,
                gap_description="No tables detected in PDF",
            )

        full_text = "\n\n".join(tables_text)
        confidence = min(0.90, 0.5 + len(tables_text) * 0.1)

        return ExtractionResult(
            text=full_text, method="TABLE_EXTRACTION",
            confidence=confidence, pages_extracted=total_pages,
            total_pages=total_pages,
        )
    except Exception as e:
        logger.debug("Table extraction failed: %s", e)
        return ExtractionResult(
            method="TABLE_EXTRACTION", confidence=0.0,
            gap_description=f"Table extraction error: {e}",
        )


# ── Strategy 2: Tesseract OCR ──────────────────────────────────────────────

def _strategy_ocr_tesseract(filepath: str, lang: str = "eng+por") -> ExtractionResult:
    """Render PDF pages as images and apply Tesseract OCR."""
    try:
        doc = fitz.open(filepath)
        try:
            total_pages = len(doc)
            page_texts = []
            all_confidences = []

            for page in doc:
                mat = fitz.Matrix(_OCR_DPI / 72, _OCR_DPI / 72)
                pixmap = page.get_pixmap(matrix=mat)
                img = Image.open(io.BytesIO(pixmap.tobytes("png")))

                ocr_data = pytesseract.image_to_data(
                    img, lang=lang, output_type=pytesseract.Output.DICT,
                )
                words = []
                for i, word in enumerate(ocr_data["text"]):
                    conf = int(ocr_data["conf"][i])
                    if conf > 20 and word.strip():
                        words.append(word)
                        all_confidences.append(conf)
                page_texts.append(" ".join(words))
        finally:
            doc.close()
        full_text = "\n\n".join(t for t in page_texts if t.strip())

        if not full_text.strip():
            return ExtractionResult(
                method="OCR_TESSERACT", confidence=0.0,
                total_pages=total_pages,
                gap_description="OCR returned no text — document may be corrupt",
            )

        avg_conf = sum(all_confidences) / len(all_confidences) / 100 if all_confidences else 0.0

        return ExtractionResult(
            text=full_text, method="OCR_TESSERACT",
            confidence=min(avg_conf, 0.85),
            pages_extracted=len([t for t in page_texts if t.strip()]),
            total_pages=total_pages,
            warnings=["OCR applied — some text may be inaccurate"],
        )
    except Exception as e:
        logger.debug("Tesseract OCR failed: %s", e)
        return ExtractionResult(
            method="OCR_TESSERACT", confidence=0.0,
            gap_description=f"Tesseract error: {e}",
        )


# ── Strategy 3: LLM vision ─────────────────────────────────────────────────

# ── R-F1666: sync fitz helpers — dispatched via asyncio.to_thread so PDF
# parsing / page rendering (fitz.open, get_text, get_pixmap) NEVER blocks the
# event loop. Rendering a large/scanned PDF ON the loop froze the whole brain —
# the doc read AND every other request (WA brain-fetch, /api/aria/health,
# state_store) timed out together (2026-06-18 contract-review outage; wedge
# stack: main thread blocked in read_document). The text/table/OCR strategies
# were already offloaded (~lines 193/205/222); the VISION path was the missed
# one. All blocking fitz work now lives in these sync helpers, off the loop.

def _fitz_page_count(filepath: str) -> int:
    doc = fitz.open(filepath)
    try:
        return len(doc)
    finally:
        doc.close()


def _fitz_vision_pdf_summaries_sync(filepath: str) -> tuple[list[str], int, int]:
    """Sync page-text extraction for _strategy_vision_pdf (runs in a thread)."""
    doc = fitz.open(filepath)
    try:
        total_pages = len(doc)
        pages_to_process = min(total_pages, _VISION_MAX_PAGES)
        page_summaries: list[str] = []
        for i in range(pages_to_process):
            page = doc[i]
            text = page.get_text()
            if text.strip():
                page_summaries.append(f"[Page {i+1}]: {text[:2000]}")
            else:
                page_summaries.append(f"[Page {i+1}]: (image-only, no text layer)")
        return page_summaries, total_pages, pages_to_process
    finally:
        doc.close()


def _fitz_chunk_summaries_sync(
    filepath: str, start_page: int, end_page: int,
) -> list[str]:
    """Sync page-text + image-render for _vision_single_chunk (runs in a thread).
    get_pixmap rasterisation is the heaviest CPU op in the doc path — keeping it
    off the loop is the core of R-F1666."""
    doc = fitz.open(filepath)
    try:
        page_summaries: list[str] = []
        for page_num in range(start_page, end_page):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                page_summaries.append(f"[PAGE {page_num + 1}]:\n{text[:3000]}")
            else:
                mat = fitz.Matrix(150 / 72, 150 / 72)
                pixmap = page.get_pixmap(matrix=mat)
                img_bytes = pixmap.tobytes("jpeg")
                if len(img_bytes) > 4_500_000:
                    mat = fitz.Matrix(100 / 72, 100 / 72)
                    pixmap = page.get_pixmap(matrix=mat)
                    img_bytes = pixmap.tobytes("jpeg")
                page_summaries.append(
                    f"[PAGE {page_num + 1}]: (image-only, {len(img_bytes)} bytes, no text layer)"
                )
        return page_summaries
    finally:
        doc.close()


async def _strategy_vision_pdf(
    filepath: str,
    llm: "LLMProvider",
    query: str = "",
) -> ExtractionResult:
    """Render PDF pages as images and send to LLM vision API.

    Note: This uses the raw LLM provider. If the provider doesn't support
    vision/multimodal, this will fail gracefully.
    """
    try:
        # R-F1666: fitz parsing/get_text off the event loop (was blocking the
        # whole brain on large/scanned PDFs).
        page_summaries, total_pages, pages_to_process = await asyncio.to_thread(
            _fitz_vision_pdf_summaries_sync, filepath,
        )

        if not page_summaries:
            return ExtractionResult(
                method="VISION_MODEL", confidence=0.0,
                gap_description="No pages rendered",
            )

        focus = f" Focus on: {query}" if query else ""
        prompt = (
            f"Extract and reconstruct all text from these {pages_to_process} "
            f"document pages. Preserve logical reading order.{focus}\n\n"
            + "\n\n".join(page_summaries)
        )

        try:
            from . import cost_tracker
            with cost_tracker.feature("document_reader"):
                result = await llm.complete(
                    "You are ARIA's document extraction system. Extract all text exactly as it appears.",
                    prompt,
                    max_tokens=6000,  # R-F1311: bumped from 4096 — complex multi-page docs need more output budget
                    timeout=90.0,     # R-F1311: bumped from 60s to match longer extraction
                )
            extracted = (getattr(result, "text", "") or "").strip()
        except Exception as e:
            return ExtractionResult(
                method="VISION_MODEL", confidence=0.0,
                gap_description=f"LLM extraction failed: {e}",
            )

        if len(extracted) < 100:
            return ExtractionResult(
                method="VISION_MODEL", confidence=0.15,
                total_pages=total_pages, pages_extracted=pages_to_process,
                gap_description="LLM returned minimal text",
            )

        warnings = []
        if pages_to_process < total_pages:
            warnings.append(
                f"Extraction limited to first {pages_to_process} of {total_pages} pages"
            )

        return ExtractionResult(
            text=extracted, method="VISION_MODEL",
            confidence=0.75, pages_extracted=pages_to_process,
            total_pages=total_pages, warnings=warnings,
        )
    except Exception as e:
        logger.debug("Vision extraction failed: %s", e)
        return ExtractionResult(
            method="VISION_MODEL", confidence=0.0,
            gap_description=f"Vision extraction failed: {e}",
        )


async def _strategy_vision_image(
    filepath: str,
    llm: "LLMProvider",
    query: str = "",
) -> ExtractionResult:
    """Extract text from a standalone image file using LLM."""
    try:
        focus = f" Focus on: {query}" if query else ""
        prompt = (
            f"Extract all text from the image at: {os.path.basename(filepath)}.{focus}\n"
            f"If the image contains a business card, receipt, or document, "
            f"extract every field and value."
        )
        from . import cost_tracker
        with cost_tracker.feature("document_reader"):
            result = await llm.complete(
                "You are ARIA's OCR system. Extract all text exactly as it appears.",
                prompt, max_tokens=4096, timeout=60.0,  # R-F1311: bumped from 2000/30s — complex images need more budget
            )
        text = (getattr(result, "text", "") or "").strip()
        return ExtractionResult(
            text=text, method="VISION_IMAGE",
            confidence=0.75 if len(text) > 50 else 0.30,
            pages_extracted=1, total_pages=1,
        )
    except Exception as e:
        return ExtractionResult(
            method="VISION_IMAGE", confidence=0.0,
            gap_description=f"Image extraction failed: {e}",
        )


# ── Chunked vision processing (large PDFs) ────────────────────────────────

async def _read_pdf_chunked(
    source: str,
    llm: "LLMProvider",
    query: str = "",
    max_pages: int = 300,
) -> ExtractionResult:
    """Process a large PDF in VISION_CHUNK_SIZE-page chunks via LLM vision.

    For documents <= VISION_LARGE_DOC_THRESHOLD pages, delegates to the
    existing single-pass ``_strategy_vision_pdf``.  For larger documents,
    renders each chunk to JPEG images, sends them individually to the LLM,
    and merges the results.  Individual chunk failures are recorded but do
    not abort the overall extraction.

    When *query* is provided and the document exceeds the large-doc
    threshold, a query-focused summary pass is appended.
    """
    if not PYMUPDF_AVAILABLE:
        return ExtractionResult(
            method="VISION_CHUNKED",
            confidence=0.0,
            gap_description="PyMuPDF not installed — cannot render pages",
        )

    try:
        # R-F1666: page-count probe off the event loop.
        total_pages = await asyncio.to_thread(_fitz_page_count, source)
    except Exception as e:
        return ExtractionResult(
            method="VISION_CHUNKED",
            confidence=0.0,
            gap_description=f"Could not open PDF: {e}",
        )

    # Small document — delegate to single-pass vision
    if total_pages <= VISION_LARGE_DOC_THRESHOLD:
        return await _strategy_vision_pdf(source, llm, query)

    pages_to_process = min(total_pages, max_pages)
    chunk_size = VISION_CHUNK_SIZE
    num_chunks = math.ceil(pages_to_process / chunk_size)

    logger.info(
        "Chunked vision processing: %d pages in %d chunks of %d",
        pages_to_process, num_chunks, chunk_size,
    )

    chunk_texts: list[str] = []
    chunk_failures_list: list[str] = []
    total_extracted_pages = 0
    confidence_scores: list[float] = []

    # R-F1672 (surgical): process chunks CONCURRENTLY (bounded) instead of
    # strictly sequentially. The old serial loop ran N vision-LLM calls one
    # after another (up to 30 chunks × ~60-120s = 30-60 min for a max-size doc)
    # — far beyond the WhatsApp 15-min poll window, so large/scanned documents
    # NEVER delivered. Running up to ARIA_VISION_CHUNK_CONCURRENCY (default 4)
    # chunks at once cuts wall-time to ~ceil(N/conc) × per-chunk, while PRESERVING
    # page order (results re-sorted by chunk index before merge), failure
    # handling, and counts. Each chunk's fitz render is already off the event
    # loop (R-F1666); the semaphore bounds concurrent LLM calls + memory.
    _chunk_conc = int(os.getenv("ARIA_VISION_CHUNK_CONCURRENCY", "4"))
    _chunk_sem = asyncio.Semaphore(max(1, _chunk_conc))

    async def _run_chunk(_idx: int):
        _sp = _idx * chunk_size
        _ep = min(_sp + chunk_size, pages_to_process)
        async with _chunk_sem:
            logger.info("Processing chunk %d/%d: pages %d-%d",
                        _idx + 1, num_chunks, _sp + 1, _ep)
            _cr = await _vision_single_chunk(
                filepath=source, start_page=_sp, end_page=_ep,
                total_pages=total_pages, llm=llm, query=query,
            )
        return _idx, _sp, _ep, _cr

    _chunk_results = await asyncio.gather(
        *[_run_chunk(i) for i in range(num_chunks)], return_exceptions=True,
    )
    # Re-assemble in PAGE ORDER — gather may complete out of order.
    _ok = sorted(
        (r for r in _chunk_results if not isinstance(r, Exception)),
        key=lambda r: r[0],
    )
    for _idx, _sp, _ep, chunk_result in _ok:
        if chunk_result.confidence > 0.10 and chunk_result.text.strip():
            chunk_texts.append(chunk_result.text)
            total_extracted_pages += chunk_result.pages_extracted
            confidence_scores.append(chunk_result.confidence)
            logger.info("Chunk %d OK: %d pages, %.0f%% confidence",
                        _idx + 1, chunk_result.pages_extracted,
                        chunk_result.confidence * 100)
        else:
            chunk_failures_list.append(f"{_sp + 1}-{_ep}")
            logger.warning("Chunk %d failed: %s", _idx + 1, chunk_result.gap_description)
    for r in _chunk_results:
        if isinstance(r, Exception):
            chunk_failures_list.append("exception")
            logger.warning("Chunk raised during parallel extraction: %s", r)

    if not chunk_texts:
        return ExtractionResult(
            method="VISION_CHUNKED",
            confidence=0.0,
            total_pages=total_pages,
            chunks_processed=0,
            chunks_total=num_chunks,
            chunk_failures=chunk_failures_list,
            gap_description=(
                f"All {num_chunks} chunks failed. "
                f"Document may be corrupt or unreadable."
            ),
        )

    merged_text = "\n\n".join(chunk_texts)
    avg_confidence = (
        sum(confidence_scores) / len(confidence_scores)
        if confidence_scores else 0.0
    )

    warnings: list[str] = []
    if chunk_failures_list:
        warnings.append(
            f"Extraction incomplete: pages {', '.join(chunk_failures_list)} "
            f"could not be processed."
        )
    if pages_to_process < total_pages:
        warnings.append(
            f"Document capped at {pages_to_process} of {total_pages} pages "
            f"(VISION_MAX_TOTAL_PAGES limit)."
        )

    # Query-focused summary pass for very large documents
    if pages_to_process > VISION_LARGE_DOC_THRESHOLD and query:
        logger.info(
            "Large document (%d pages) with query — running summary pass",
            pages_to_process,
        )
        summary = await _summarise_large_document(merged_text, query, llm)
        if summary:
            merged_text = (
                f"DOCUMENT SUMMARY (query-focused):\n{summary}\n\n"
                f"{'─' * 60}\n"
                f"FULL EXTRACTED TEXT ({pages_to_process} pages):\n\n"
                f"{merged_text}"
            )

    logger.info(
        "Chunked extraction complete: %d pages extracted, %d chunks failed, "
        "%.0f%% average confidence",
        total_extracted_pages, len(chunk_failures_list), avg_confidence * 100,
    )

    return ExtractionResult(
        text=merged_text,
        method="VISION_CHUNKED",
        confidence=avg_confidence,
        pages_extracted=total_extracted_pages,
        total_pages=total_pages,
        warnings=warnings,
        chunks_processed=len(chunk_texts),
        chunks_total=num_chunks,
        chunk_failures=chunk_failures_list,
    )


async def _vision_single_chunk(
    filepath: str,
    start_page: int,
    end_page: int,
    total_pages: int,
    llm: "LLMProvider",
    query: str = "",
) -> ExtractionResult:
    """Render a range of PDF pages to JPEG and extract text via LLM vision.

    Returns an ``ExtractionResult`` for the specified page range.  Uses the
    existing LLM provider abstraction (``llm.complete``) rather than a
    direct Anthropic client call.
    """
    try:
        # R-F1666: fitz get_text + get_pixmap rasterisation off the event loop —
        # this rendering was the heaviest on-loop block (froze the brain on the
        # 2026-06-18 contract review).
        page_summaries = await asyncio.to_thread(
            _fitz_chunk_summaries_sync, filepath, start_page, end_page,
        )

        if not page_summaries:
            return ExtractionResult(
                method="VISION_CHUNKED",
                confidence=0.0,
                gap_description=f"No pages rendered for chunk {start_page + 1}-{end_page}",
            )

        chunk_label = f"pages {start_page + 1}-{end_page}"
        focus = f" Focus on: {query}" if query else ""
        prompt = (
            f"Extract ALL text from these {chunk_label} of a {total_pages}-page document "
            f"exactly as it appears. Include all paragraphs, headings, tables, captions, "
            f"footnotes, and any text content. Preserve logical reading order.{focus}\n\n"
            + "\n\n".join(page_summaries)
        )

        from . import cost_tracker
        with cost_tracker.feature("document_reader"):
            result = await llm.complete(
                "You are ARIA's document extraction system. Extract all text exactly as it appears.",
                prompt,
                max_tokens=6000,  # R-F1311: bumped from 4096
                timeout=120.0,   # R-F1311: bumped from 90s
            )

        extracted = (getattr(result, "text", "") or "").strip()
        pages_extracted = end_page - start_page

        return ExtractionResult(
            text=extracted,
            method="VISION_CHUNKED",
            confidence=0.80 if len(extracted) > 100 else 0.20,
            pages_extracted=pages_extracted,
            total_pages=total_pages,
        )
    except Exception as e:
        logger.warning("Chunk %d-%d error: %s", start_page + 1, end_page, e)
        return ExtractionResult(
            method="VISION_CHUNKED",
            confidence=0.0,
            gap_description=f"Chunk {start_page + 1}-{end_page} error: {e}",
        )


async def _summarise_large_document(
    full_text: str,
    query: str,
    llm: "LLMProvider",
) -> str:
    """Produce a query-focused summary of merged chunk text.

    Only intended for documents exceeding ``VISION_LARGE_DOC_THRESHOLD``
    pages where a *query* was provided.  Samples the beginning, middle,
    and end of the text to stay within token limits.
    """
    text_len = len(full_text)
    sample = full_text[:3000]
    if text_len > 6000:
        mid = text_len // 2
        sample += "\n\n[...middle section...]\n\n" + full_text[mid:mid + 2000]
    if text_len > 4000:
        sample += "\n\n[...final section...]\n\n" + full_text[-2000:]

    prompt = (
        f"This is extracted text from a large document.\n"
        f"Query: {query}\n\n"
        f"Document extract:\n{sample}\n\n"
        f"Produce a structured summary covering:\n"
        f"1. Document type and purpose\n"
        f"2. Key parties, entities, and dates\n"
        f"3. Most relevant content for the query: {query}\n"
        f"4. Key figures, provisions, or findings\n"
        f"5. Any red flags or compliance concerns\n"
        f"Maximum 600 words. Concise and precise."
    )

    try:
        from . import cost_tracker
        with cost_tracker.feature("document_reader"):
            result = await llm.complete(
                "You are ARIA's document analysis system. Produce concise structured summaries.",
                prompt,
                max_tokens=1000,
                timeout=30.0,
            )
        return (getattr(result, "text", "") or "").strip()
    except Exception as e:
        logger.warning("Summary pass failed: %s", e)
        return ""


# ── Strategy 4: Online search ──────────────────────────────────────────────

def _strategy_find_online(filepath: str, query: str = "") -> ExtractionResult:
    """Search for an accessible version of the document online."""
    try:
        metadata = _extract_pdf_metadata(filepath)
        title = metadata.get("title", "")
        if not title:
            return ExtractionResult(
                method="ONLINE_SEARCH", confidence=0.0,
                gap_description="No metadata for online search",
            )

        return ExtractionResult(
            text=f"[Document identified: '{title}'. Search online for accessible version.]",
            method="ONLINE_SEARCH", confidence=0.20,
            gap_description=f"Original unreadable. Title: {title}",
        )
    except Exception as e:
        return ExtractionResult(
            method="ONLINE_SEARCH", confidence=0.0,
            gap_description=f"Online search failed: {e}",
        )


# ── Non-PDF readers ─────────────────────────────────────────────────────────

def _read_plaintext(filepath: str) -> ExtractionResult:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(filepath, "r", encoding=encoding) as f:
                text = f.read()
            return ExtractionResult(
                text=text, method="PLAINTEXT", confidence=0.99,
                pages_extracted=1, total_pages=1,
            )
        except UnicodeDecodeError:
            continue
    return ExtractionResult(method="PLAINTEXT", confidence=0.0,
                            gap_description="Encoding unrecognised")


def _read_html(filepath: str) -> ExtractionResult:
    try:
        import trafilatura
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()
        text = trafilatura.extract(html) or ""
        if len(text) > 100:
            return ExtractionResult(
                text=text, method="HTML_TRAFILATURA", confidence=0.90,
                pages_extracted=1, total_pages=1,
            )
    except ImportError:
        pass
    except Exception as e:
        logger.debug("trafilatura failed: %s", e)

    try:
        from bs4 import BeautifulSoup
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()
        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text(separator=" ", strip=True)
        return ExtractionResult(
            text=text, method="HTML_BS4",
            confidence=0.75 if len(text) > 100 else 0.20,
            pages_extracted=1, total_pages=1,
        )
    except Exception as e:
        return ExtractionResult(method="HTML", confidence=0.0,
                                gap_description=f"HTML extraction failed: {e}")


def _read_docx(filepath: str) -> ExtractionResult:
    # R-F946 (2026-05-27) — extract the body IN DOCUMENT ORDER. The previous
    # impl read ALL paragraphs, then APPENDED all table cells at the end, so a
    # contract with interleaved tables came out jumbled: the live Korvera review
    # stopped mid-sentence at "Clause 5.4(b) … cancellation of th" while later
    # clauses were dumped out of order, and text-boxes / content-controls were
    # missed entirely. Walk word/document.xml (+ headers, footers, foot/endnotes)
    # converting structural boundaries (paragraph, row, tab, break) to whitespace
    # and keeping the run text in place — this captures tables, text-boxes and
    # SDT content in reading order so ARIA sees the WHOLE document coherently.
    #
    # R-F1437: strip tracked-change DELETIONS entirely (w:del blocks) so
    # struck-through text never appears as live contract terms. Keep w:ins
    # insertions — those are the live/new terms the author accepted.
    import re as _re
    try:
        import zipfile
        import html as _html
        texts: list[str] = []
        with zipfile.ZipFile(filepath) as zf:
            names = set(zf.namelist())
            ordered = (
                ["word/document.xml"]
                + sorted(n for n in names if _re.match(r"word/header\d*\.xml$", n))
                + sorted(n for n in names if _re.match(r"word/footnotes\.xml$", n))
                + sorted(n for n in names if _re.match(r"word/endnotes\.xml$", n))
                + sorted(n for n in names if _re.match(r"word/footer\d*\.xml$", n))
            )
            for n in ordered:
                if n not in names:
                    continue
                xml = zf.read(n).decode("utf-8", errors="ignore")
                # R-F1437: strip tracked-change DELETIONS entirely.
                # In OOXML, deleted text lives in <w:del><w:r><w:delText>...</w:delText></w:r></w:del>.
                # Removing the entire w:del block ensures struck-through text never
                # appears as live contract terms. Keep w:ins insertions — those are
                # the live/new terms the author accepted (they contain normal w:r runs).
                # R-F1757: use [^<] instead of .*? to avoid catastrophic backtracking
                # on large documents with many tracked changes. The old .*? with
                # re.DOTALL could hang indefinitely on multi-MB XML.
                xml = _re.sub(r"<w:del\b[^>]*>[^<]*(?:<(?!/w:del>)[^<]*)*</w:del>", "", xml)
                # Drop field instruction codes (TOC/HYPERLINK/REF) — they are not
                # body text and would inject noise.
                xml = _re.sub(r"<w:instrText[^>]*>[^<]*(?:<(?!/w:instrText>)[^<]*)*</w:instrText>", "", xml)
                # Structural boundaries → whitespace, BEFORE stripping tags.
                xml = _re.sub(r"<w:p\b[^>]*>", "\n", xml)      # paragraph
                xml = _re.sub(r"<w:tr\b[^>]*>", "\n", xml)     # table row
                xml = _re.sub(r"<w:tab\b[^>]*/?>", "\t", xml)  # tab
                xml = _re.sub(r"<w:br\b[^>]*/?>", "\n", xml)   # line break
                texts.append(_re.sub(r"<[^>]+>", "", xml))     # strip remaining tags; run text survives
        text = _html.unescape("".join(texts))
        text = _re.sub(r"[ \t]+\n", "\n", text)
        text = _re.sub(r"\n{3,}", "\n\n", text).strip()
        if len(text) > 100:
            return ExtractionResult(
                text=text, method="DOCX", confidence=0.95,
                pages_extracted=1, total_pages=1,
            )
        # raw parse yielded almost nothing — fall through to python-docx
    except Exception as e:
        logger.debug("DOCX in-order XML extraction failed, falling back to python-docx: %s", e)

    # Fallback: python-docx (paragraphs + tables). Ordering is less faithful but
    # it still recovers the body if the zip/xml parse path failed.
    # R-F1437: python-docx includes deleted tracked-change text by default.
    # We filter it out by re-parsing the XML and stripping w:del blocks.
    try:
        from docx import Document
        doc = Document(filepath)
        # R-F1437: filter deleted tracked-change text from python-docx output.
        # python-docx doesn't expose tracked changes natively, so we re-read
        # the XML to strip w:del blocks, then re-extract text from the clean XML.
        try:
            import zipfile as _zf
            _clean_xml = ""
            with _zf.ZipFile(filepath) as _z:
                if "word/document.xml" in _z.namelist():
                    _raw = _z.read("word/document.xml").decode("utf-8", errors="ignore")
                    _clean_xml = _re.sub(
                        # R-F1757: safe [^<]* pattern (no catastrophic backtracking)
                        r"<w:del\b[^>]*>[^<]*(?:<(?!/w:del>)[^<]*)*</w:del>", "", _raw,
                    )
            if _clean_xml:
                # Re-parse with python-docx from the clean XML
                import io as _io
                import zipfile as _zf2
                _buf = _io.BytesIO()
                with _zf2.ZipFile(filepath, "r") as _src:
                    with _zf2.ZipFile(_buf, "w") as _dst:
                        for _item in _src.infolist():
                            _data = _src.read(_item.filename)
                            if _item.filename == "word/document.xml":
                                _data = _clean_xml.encode("utf-8")
                            _dst.writestr(_item, _data)
                _buf.seek(0)
                doc = Document(_buf)
        except Exception:
            pass  # Fall through to original doc if XML manipulation fails
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        return ExtractionResult(
            text=text, method="DOCX",
            confidence=0.95 if len(text) > 100 else 0.20,
            pages_extracted=1, total_pages=1,
        )
    except ImportError:
        return ExtractionResult(method="DOCX", confidence=0.0,
                                gap_description="python-docx not installed")
    except Exception as e:
        return ExtractionResult(method="DOCX", confidence=0.0,
                                gap_description=f"DOCX error: {e}")


# ── R-F1104: .xlsx spreadsheet reader ─────────────────────────────────────


def _read_xlsx(filepath: str) -> ExtractionResult:
    """Read a .xlsx spreadsheet and return all sheets as structured text.

    Each sheet is rendered as:
      === Sheet: <name> ===
      | Col A | Col B | Col C |
      | val1  | val2  | val3  |

    Returns ExtractionResult with method="XLSX".
    """
    if not OPENPYXL_AVAILABLE:
        return ExtractionResult(
            method="XLSX", confidence=0.0,
            gap_description="openpyxl not installed — install with: pip install openpyxl",
        )
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        parts: list[str] = []
        total_rows = 0
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                # Filter fully-empty rows
                cleaned = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cleaned):
                    parts.append(" | ".join(cleaned))
                    total_rows += 1
            parts.append("")  # blank line between sheets
        wb.close()
        text = "\n".join(parts)
        if len(text) > 50:
            return ExtractionResult(
                text=text, method="XLSX", confidence=0.90,
                pages_extracted=total_rows, total_pages=total_rows,
            )
        return ExtractionResult(
            method="XLSX", confidence=0.10,
            gap_description="XLSX file appears empty or all-blank",
        )
    except Exception as e:
        logger.debug("XLSX read failed for %s: %s", filepath, e)
        return ExtractionResult(
            method="XLSX", confidence=0.0,
            gap_description=f"XLSX error: {e}",
        )


# ── Utilities ───────────────────────────────────────────────────────────────

async def _resolve_source(source: str) -> str | None:
    """Download URL to temp file or return file path as-is."""
    if source.startswith(("http://", "https://")):
        try:
            import httpx
            from . import url_safety as _us
            async with httpx.AsyncClient(timeout=_REQUEST_TIMEOUT) as client:
                # R-F1814 (audit C2): SSRF-guarded fetch (validates url + every
                # redirect hop; follow_redirects off). `source` is user-controlled.
                resp = await _us.safe_get(client, source)
                resp.raise_for_status()
                ct = resp.headers.get("content-type", "")
                suffix = ".pdf"
                if "html" in ct:
                    suffix = ".html"
                elif "plain" in ct:
                    suffix = ".txt"
                with tempfile.NamedTemporaryFile(
                    suffix=suffix, delete=False, prefix="aria_doc_"
                ) as f:
                    f.write(resp.content)
                    return f.name
        except ValueError as _ssrf:  # ssrf_blocked:<reason> from safe_get
            logger.warning("document_reader: blocked SSRF-unsafe source %r (%s)", str(source)[:80], _ssrf)
            return None
        except Exception as e:
            logger.debug("URL download failed for %s: %s", source[:80], e)
            return None

    # R-F1814 (audit C2): local-file branch — restrict to the system temp dir
    # (where uploads land via NamedTemporaryFile) to prevent arbitrary-file read.
    # `source` is user-controlled; '/etc/passwd', 'C:\\Windows\\...' must be rejected.
    _abs = os.path.realpath(source)
    _tmp_root = os.path.realpath(tempfile.gettempdir())
    if _abs.startswith(_tmp_root + os.sep) and os.path.exists(_abs):
        return _abs
    logger.warning("document_reader: rejected non-tempdir local source %r", str(source)[:80])
    return None


def _extract_pdf_metadata(filepath: str) -> dict:
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(filepath) as pdf:
                return pdf.metadata or {}
        except Exception:
            pass
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(filepath)
            try:
                meta = doc.metadata
            finally:
                doc.close()
            return meta or {}
        except Exception:
            pass
    return {}


def _resolve_ocr_language(hint: str) -> str:
    lang_map = {
        "por": "eng+por", "pt": "eng+por",
        "fra": "eng+fra", "fr": "eng+fra",
        "ara": "eng+ara", "ar": "eng+ara",
        "": _OCR_LANGUAGES,
    }
    return lang_map.get(hint.lower(), _OCR_LANGUAGES)


# ── Contract intelligence (SITCL detection) ─────────────────────────────────

SITCL_TRIGGERS = [
    "brokering", "broker", "arrange", "arranging",
    "military goods", "controlled goods", "controlled items",
    "dual-use", "munitions", "arms", "weapons",
    "export licence", "export license", "end user",
    "end-user certificate", "euc",
]

REQUIRED_CONTRACT_CLAUSES = [
    "parties", "jurisdiction", "governing law", "scope of work",
    "payment terms", "exclusivity", "intellectual property",
    "termination", "liability cap", "representations and warranties",
    "compliance", "force majeure", "anti-assignment", "entire agreement",
]


@fail_wire(module="document_reader", gap_type="file_parse")
async def analyse_contract(
    source: str,
    llm: "LLMProvider | None" = None,
    market: str = "",
) -> dict:
    """Read and analyse a contract document for SITCL triggers and missing clauses."""
    extraction = await read_document(source, llm=llm, query="legal contract terms")

    if not extraction.is_usable:
        return {
            "status": "UNREADABLE",
            "extraction": extraction.summary,
            "analysis": None,
        }

    text_lower = extraction.text.lower()

    missing = [c for c in REQUIRED_CONTRACT_CLAUSES if c not in text_lower]
    sitcl_found = [kw for kw in SITCL_TRIGGERS if kw in text_lower]


    # R-F996 — wire to brain
    wire_success(
        module="document_reader",
        summary="Document reading",
        source_id="document_reader:R-F996",
    )
    return {
        "status": "ANALYSED",
        "extraction_method": extraction.method,
        "extraction_confidence": extraction.confidence,
        "missing_clauses": missing,
        "sitcl_required": len(sitcl_found) > 0,
        "sitcl_triggers": sitcl_found,
        "reference": f"ARK-READ-{datetime.now().strftime('%Y%m%d%H%M%S')}",
    }


@fail_wire(module="document_reader", gap_type="file_parse")
async def analyse_contract_deep(
    source: str,
    llm: "LLMProvider | None" = None,
    market: str = "",
    comparison_source: str | None = None,
) -> dict:
    """Deep LLM-powered contract analysis with optional redline comparison.

    Unlike analyse_contract() which does keyword-based SITCL/missing-clause
    detection, this function uses the LLM to produce a full commercial
    analysis: risk assessment, party-benefit analysis, negotiation
    recommendations, and (when comparison_source is provided) a redline
    comparison between two versions of a contract.

    Args:
        source: Path or URL to the primary contract document (newer version
            in a redline comparison).
        llm: LLM provider for the analysis.
        market: Optional market context (e.g. "defence", "Peru").
        comparison_source: Optional path/URL to the older version for
            redline comparison.

    Returns:
        dict with keys:
          status: "ANALYSED" | "UNREADABLE" | "ERROR"
          analysis: LLM-generated analysis text
          redline_changes: list of detected changes (if comparison provided)
          risk_flags: list of high/medium/low risk items
          negotiation_points: list of recommended negotiation items
          extraction_method: how the document was read
    """
    # Read the primary document
    extraction = await read_document(source, llm=llm, query="legal contract terms")
    if not extraction.is_usable:
        return {
            "status": "UNREADABLE",
            "extraction": extraction.summary,
            "analysis": None,
        }

    # Read comparison document if provided
    comparison_text = None
    if comparison_source:
        comp_extraction = await read_document(
            comparison_source, llm=llm, query="legal contract terms comparison"
        )
        if comp_extraction.is_usable:
            comparison_text = comp_extraction.text

    # Build the analysis prompt
    doc_text = extraction.text
    # Truncate to avoid blowing the LLM context window
    max_doc_chars = 25000
    if len(doc_text) > max_doc_chars:
        doc_text = doc_text[:max_doc_chars] + "\n\n[... document truncated ...]"

    market_context = f" The contract relates to the {market} market." if market else ""

    if comparison_text:
        if len(comparison_text) > max_doc_chars:
            comparison_text = comparison_text[:max_doc_chars] + "\n\n[... document truncated ...]"
        prompt = (
            f"You are a senior commercial contracts analyst. Analyse the following "
            f"contract REDLINE COMPARISON between two versions (OLDER vs NEWER).{market_context}\n\n"
            f"--- OLDER VERSION ---\n{comparison_text}\n\n"
            f"--- NEWER VERSION ---\n{doc_text}\n\n"
            f"Provide a structured analysis covering:\n"
            f"1. KEY CHANGES: List each material change between versions, which party "
            f"benefits, and the commercial impact.\n"
            f"2. RISK ASSESSMENT: For each change and for the overall agreement, rate "
            f"risk as HIGH/MEDIUM/LOW and explain why.\n"
            f"3. NEGOTIATION RECOMMENDATIONS: Specific items the receiving party should "
            f"push back on, with suggested language or approach.\n"
            f"4. PARTY ANALYSIS: Which party does each clause favour? Is the overall "
            f"agreement balanced or one-sided?\n"
            f"5. STRUCTURAL ISSUES: Missing clauses, unbalanced termination provisions, "
            f"liability caps, dispute resolution concerns.\n"
            f"6. EXECUTIVE SUMMARY: 3-5 bullet points summarising the most important "
            f"things to know about this agreement."
        )
    else:
        prompt = (
            f"You are a senior commercial contracts analyst. Analyse the following "
            f"contract agreement.{market_context}\n\n"
            f"--- CONTRACT ---\n{doc_text}\n\n"
            f"Provide a structured analysis covering:\n"
            f"1. PARTIES AND SCOPE: Who are the parties, what is the territory, "
            f"what products/services are covered?\n"
            f"2. KEY TERMS: Commission structure, exclusivity, term, termination rights.\n"
            f"3. RISK ASSESSMENT: For each major clause, rate risk as HIGH/MEDIUM/LOW "
            f"and explain why. Consider: termination provisions, liability caps, "
            f"indemnities, governing law, dispute resolution, non-compete, IP rights.\n"
            f"4. PARTY ANALYSIS: Which party does each clause favour? Is the overall "
            f"agreement balanced or one-sided?\n"
            f"5. NEGOTIATION RECOMMENDATIONS: Specific items to negotiate, with "
            f"suggested approach.\n"
            f"6. STRUCTURAL ISSUES: Missing clauses, unbalanced provisions, "
            f"jurisdiction concerns.\n"
            f"7. EXECUTIVE SUMMARY: 3-5 bullet points summarising the most important "
            f"things to know."
        )

    if not llm or not llm.is_configured:
        return {
            "status": "ERROR",
            "error": "LLM not configured — cannot perform deep analysis",
            "extraction_method": extraction.method,
        }

    try:
        result = await llm.complete(prompt, max_tokens=4000, temperature=0.3)
        analysis_text = result.content if result and result.content else "Analysis failed to generate."
    except Exception as e:
        logger.error("analyse_contract_deep LLM call failed: %s", e)
        return {
            "status": "ERROR",
            "error": f"LLM analysis failed: {e}",
            "extraction_method": extraction.method,
        }

    # R-F996 — wire to brain
    wire_success(
        module="document_reader",
        summary="Deep contract analysis",
        source_id="document_reader:R-F2078",
    )

    return {
        "status": "ANALYSED",
        "analysis": analysis_text,
        "is_redline": comparison_text is not None,
        "extraction_method": extraction.method,
        "extraction_confidence": extraction.confidence,
        "reference": f"ARK-CONTRACT-{datetime.now().strftime('%Y%m%d%H%M%S')}",
    }

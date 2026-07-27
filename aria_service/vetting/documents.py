"""R-F3144/R-F3145 — evidence intake and document extraction for vetting.

── Why dd_evidence_store and not a new bucket (R-F3144) ──────────────────
The module README proposed presigned uploads to Tigris. This repo already
has `intel/dd_evidence_store.py` (R-F3083): content-addressed artifact
retention, sha256 verification on append, an explicit tenant boundary, and
idempotent re-append. That is exactly what BS 7858's record-keeping
expectation asks for — a retained copy, who examined it, when — and §6 puts
the burden of proof on any new third party. So evidence lands there.

The evidence vocabulary already models this case honestly: an applicant's
payslip is `SourceAuthority.USER_SUPPLIED`, not a primary-official record.
Recording it as anything stronger would overstate the evidence at the point
it enters the file, which is the hardest place to correct later.

── What extraction may and may not decide (R-F3145) ──────────────────────
The model reads documents. It does not grade people, and it does not get to
mark its own homework. Every extraction lands as a PROPOSAL:

  * A confident classification sets doc_type and the covered dates.
  * Anything below the confidence floor, or any authenticity concern, is
    recorded as an authenticity_flag and the document is left for a human —
    it is never silently accepted.
  * An LLM failure (unreachable, timeout, invalid output) yields NO
    extraction at all. The document is still stored and still counted as
    present; it simply carries `extraction_unavailable`. A failed read must
    never look like a document that was read and found unremarkable.

That last rule is the whole point. Everything the employer relies on — gaps,
clocks, thresholds — is computed deterministically by rules.py from what is
on the file. Extraction only proposes what a document IS.
"""

from __future__ import annotations

import hashlib
import json
import logging
import uuid
from datetime import UTC, date, datetime
from typing import Any

from .models import CareerEntryType, DocumentType, UploadedDocument

_log = logging.getLogger(__name__)

# Below this, a classification is a suggestion for a human, not a fact.
CONFIDENCE_FLOOR = 0.70

MAX_DOCUMENT_BYTES = 16 * 1024 * 1024

_EXTRACTION_SYSTEM_PROMPT = """\
You classify documents submitted for pre-employment screening.

Return ONLY what the document itself shows. Rules:
- If you cannot tell what the document is, use doc_type "OTHER" and a low
  confidence. Guessing is worse than admitting uncertainty here: a
  misclassified document silently changes which evidence a screening file
  appears to hold.
- covers_from / covers_to are the period the document EVIDENCES (e.g. the
  pay period on a payslip, the tax year on a P60), not the date it was
  printed. Use null if the document does not state a period.
- issuer is the organisation that produced the document, verbatim.
- Report anything that would make a careful human look twice in
  authenticity_concerns: inconsistent fonts or alignment, dates that
  contradict each other, evidence of digital editing, a mismatched logo,
  an altered total. Do not speculate about intent; describe what you see.
- You are NOT assessing the person. Never infer suitability, character or
  risk. Classify the document only.
- declared_periods: ONLY for an APPLICATION_FORM, and only for periods the
  form itself states. This is the applicant's declared history — employment,
  self-employment, unemployment, education, career breaks, time abroad. Copy
  the dates as written; do not infer, tidy or fill a gap. If a period has no
  stated end date leave end null. For any other kind of document return an
  empty list: a payslip that happens to mention other dates is not a
  declaration of history.
"""

_EXTRACTION_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "doc_type": {"type": "string",
                     "enum": [t.value for t in DocumentType]},
        "issuer": {"type": ["string", "null"]},
        "covers_from": {"type": ["string", "null"],
                        "description": "ISO date YYYY-MM-DD or null"},
        "covers_to": {"type": ["string", "null"],
                      "description": "ISO date YYYY-MM-DD or null"},
        "confidence": {"type": "number", "minimum": 0, "maximum": 1},
        "authenticity_concerns": {"type": "array", "items": {"type": "string"}},
        # R-F3274 — an application form is the applicant DECLARING a history.
        # Only that document type fills this; see the prompt. Consumed by
        # timeline.py, which files every period as UNVERIFIED (declared, not
        # evidenced) and records which document it was read from.
        "declared_periods": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "entry_type": {"type": "string",
                                   "enum": [t.value for t in CareerEntryType]},
                    "start": {"type": ["string", "null"],
                              "description": "ISO date YYYY-MM-DD"},
                    "end": {"type": ["string", "null"],
                            "description": "ISO date YYYY-MM-DD, or null if ongoing"},
                    "organisation": {"type": ["string", "null"]},
                },
                "required": ["entry_type", "start"],
            },
        },
    },
    "required": ["doc_type", "confidence"],
}


def sha256_hex(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _parse_iso_date(value: Any) -> date | None:
    if not value or not isinstance(value, str):
        return None
    try:
        return date.fromisoformat(value.strip()[:10])
    except ValueError:
        return None


# The evidence contract (dd_evidence_standard) requires evidence_id, case_id,
# subject_entity_id and source_attempt_id to be UUIDs. Vetting case ids are
# human references ("VET-2026-001") and subjects are names, so they are mapped
# into a UUID namespace with uuid5 — DETERMINISTIC, so the same case and the
# same bytes always produce the same ids, which is what keeps the store's
# de-duplication working across restarts. A uuid4 here would silently turn one
# re-uploaded payslip into two pieces of evidence.
_VETTING_NS = uuid.UUID("5f7d4a2e-9c11-4b6e-8a3f-2d6c1e0b7a94")


def _stable_uuid(*parts: str) -> str:
    return str(uuid.uuid5(_VETTING_NS, "|".join(parts)))


def build_evidence_record(
    *,
    tenant_id: str,
    case_id: str,
    evidence_id: str,
    content_hash: str,
    filename: str,
    subject_entity_id: str,
) -> dict[str, Any]:
    """One canonical evidence record for an applicant-supplied document."""
    return {
        "evidence_id": evidence_id,
        "tenant_id": tenant_id,
        "case_id": _stable_uuid("case", tenant_id, case_id),
        "case_scope_version": 1,
        "subject_entity_id": _stable_uuid(
            "subject", tenant_id, case_id, subject_entity_id),
        # Stable per upload: the same bytes uploaded twice to the same case
        # de-duplicate rather than double-count as two pieces of evidence.
        "source_attempt_id": _stable_uuid(
            "attempt", tenant_id, case_id, content_hash),
        "source_id": "vetting_applicant_upload",
        "source_authority": "user_supplied",
        "retrieval_outcome": "success",
        "retrieved_at": datetime.now(UTC).isoformat(),
        "request_fingerprint": sha256_hex(
            f"{tenant_id}|{case_id}|{filename}|{content_hash}".encode()),
        "licence_policy_id": "vetting_applicant_submission",
        "access_method": "applicant_upload",
        "adapter_version": "vetting-upload-1",
        "parser_version": "vetting-extract-1",
        "content_hash": content_hash,
        # PERMITTED: the standard expects the examined copy to be retained on
        # the screening file, and retention is a documented obligation here
        # rather than an optional convenience. The contract then REQUIRES a
        # raw_artifact_uri; the store is content-addressed, so the hash IS the
        # locator.
        "snapshot_policy": "permitted",
        "raw_artifact_uri": f"vetting-evidence://{content_hash}",
        "source_url": None,
        "structured_payload": {"filename": filename},
    }


async def extract_document(
    *,
    text: str,
    filename: str,
    llm: Any = None,
) -> dict[str, Any]:
    """Classify one document. NEVER raises, and never fabricates on failure.

    Returns a dict with `available` False when no usable extraction could be
    obtained — the caller must treat that as "not yet read", never as "read
    and unremarkable".
    """
    from ..llm.structured import call_structured
    from .processors import assert_served_by_approved, resolve_vetting_processor

    if not text.strip():
        return {"available": False, "reason": "no extractable text in document"}

    # R-F3151 — NEVER the application chain. Passing llm=None here would let
    # call_structured resolve app.state.llm_provider, whose documented default
    # is "everything else on DeepSeek" — an unapproved processor with no
    # adequacy decision, carrying identity and criminal-offence data. Refusing
    # to process is always available; un-disclosing is not.
    if llm is None:
        resolution = resolve_vetting_processor()
        if not resolution.usable:
            return {"available": False,
                    "reason": f"no approved processor: {resolution.reason}"}
        llm = resolution.provider

    result = await call_structured(
        _EXTRACTION_SYSTEM_PROMPT,
        f"Filename: {filename}\n\nDocument text:\n{text[:12000]}",
        schema=_EXTRACTION_SCHEMA,
        caller="vetting.extract_document",
        max_tokens=600,
        timeout=45.0,
        llm=llm,
    )
    # Detector, not a control (see processors.py): if something reintroduces a
    # chain, the transfer already happened — but it must be LOUD, and the
    # output of an unapproved processor must not be relied on.
    if not assert_served_by_approved(getattr(result, "provider", "")):
        from ..intel.engine_wiring import wire_failure
        wire_failure(
            module="vetting",
            detail=(f"vetting extraction was served by unapproved processor "
                    f"{result.provider!r} — personal data may have left an "
                    f"approved processor boundary"),
            gap_type="data_protection_violation",
            source="vetting.extract_document",
        )
        return {"available": False,
                "reason": f"served by unapproved processor {result.provider!r}"}
    if not result.ok or not isinstance(result.data, dict):
        # call_structured already wired the failure to the brain (§21a).
        return {
            "available": False,
            "reason": f"extraction unavailable ({result.outcome})",
        }
    return {"available": True, "data": result.data}


def apply_extraction(
    *,
    document_id: str,
    evidence_id: str,
    fallback_doc_type: DocumentType,
    extraction: dict[str, Any],
) -> UploadedDocument:
    """Turn a raw extraction into an UploadedDocument, conservatively.

    The confidence floor and the authenticity flags both route to the same
    place: a human looks. Nothing here can mark a document accepted.
    """
    flags: list[str] = []
    doc_type = fallback_doc_type
    issuer: str | None = None
    covers_from: date | None = None
    covers_to: date | None = None
    confidence = 0.0

    if not extraction.get("available"):
        # A failed read is a DISCLOSED gap, not a silent pass.
        flags.append("extraction_unavailable")
        flags.append(str(extraction.get("reason", "unknown"))[:200])
        return UploadedDocument(
            document_id=document_id, doc_type=doc_type, evidence_id=evidence_id,
            extraction_confidence=0.0, authenticity_flags=flags,
        )

    data = extraction.get("data") or {}
    try:
        confidence = float(data.get("confidence", 0.0) or 0.0)
    except (TypeError, ValueError):
        confidence = 0.0

    raw_type = str(data.get("doc_type", "") or "").strip().upper()
    if confidence >= CONFIDENCE_FLOOR:
        try:
            doc_type = DocumentType(raw_type)
        except ValueError:
            flags.append(f"unrecognised_doc_type:{raw_type[:40]}")
    else:
        # Keep the caller's declared type; record the suggestion for a human
        # rather than acting on it.
        flags.append(
            f"low_confidence_classification:{raw_type[:40]}@{confidence:.2f}")

    issuer = (str(data.get("issuer")).strip()
              if data.get("issuer") else None)
    covers_from = _parse_iso_date(data.get("covers_from"))
    covers_to = _parse_iso_date(data.get("covers_to"))
    if covers_from and covers_to and covers_to < covers_from:
        # A document that evidences a period ending before it starts is not a
        # period we may quietly normalise — drop both and flag it.
        flags.append("inconsistent_coverage_dates")
        covers_from = covers_to = None

    for concern in (data.get("authenticity_concerns") or [])[:10]:
        text = str(concern).strip()
        if text:
            flags.append(f"authenticity:{text[:160]}")

    return UploadedDocument(
        document_id=document_id,
        doc_type=doc_type,
        evidence_id=evidence_id,
        covers_from=covers_from,
        covers_to=covers_to,
        issuer=issuer,
        extraction_confidence=confidence,
        authenticity_flags=flags,
    )


def needs_human_review(document: UploadedDocument) -> bool:
    """True when this document must not be counted as accepted evidence."""
    return (
        bool(document.authenticity_flags)
        or document.extraction_confidence < CONFIDENCE_FLOOR
    )


def new_document_id() -> str:
    return f"vdoc_{uuid.uuid4().hex[:16]}"


def new_evidence_id() -> str:
    # Must be a bare UUID — the evidence contract validates the format.
    return str(uuid.uuid4())


# R-F3265 — bounds. An upload is a user-facing request, so extraction must not
# be able to hold one open. A 500-page scan is capped by pages; a text-heavy
# document is capped by characters. Both are generous for a screening document
# (a passport is 1 page, a reference 1-3, a bank statement rarely over 20) and
# small enough that no single upload stalls.
_MAX_EXTRACT_PAGES = 30
_MAX_EXTRACT_CHARS = 200_000
# Below this many characters, a PDF page is treated as having no text layer —
# a scan. Not zero: a scanned page often yields a few stray glyphs from a
# letterhead, and treating that as "text found" skips the OCR that holds the
# actual content.
_OCR_TEXT_FLOOR = 24


def _extract_pdf(content: bytes) -> str:
    """PDF text, with OCR only for pages that have no text layer.

    OCR is expensive, so it is a FALLBACK, not the default: a digital PDF is
    read from its text layer, and tesseract runs only on pages that came back
    empty — which is what a scan or a phone photo of a passport looks like.
    """
    try:
        import fitz  # PyMuPDF
    except Exception:  # noqa: BLE001 — an absent dep is a disclosed gap
        _log.warning("vetting: PyMuPDF unavailable — PDF left unread")
        return ""

    parts: list[str] = []
    try:
        with fitz.open(stream=content, filetype="pdf") as doc:
            for index, page in enumerate(doc):
                if index >= _MAX_EXTRACT_PAGES:
                    break
                text = ""
                try:
                    text = page.get_text() or ""
                except Exception:  # noqa: BLE001 — one bad page is not the document
                    text = ""
                if len(text.strip()) < _OCR_TEXT_FLOOR:
                    text = _ocr_page(page) or text
                parts.append(text)
                if sum(len(p) for p in parts) >= _MAX_EXTRACT_CHARS:
                    break
    except Exception as exc:  # noqa: BLE001
        # A corrupt or encrypted PDF reads as nothing — which routes to
        # `extraction_unavailable` and a human, exactly as before.
        _log.debug("vetting: PDF extraction failed: %s", exc)
        return ""
    return "\n".join(parts)[:_MAX_EXTRACT_CHARS]


def _ocr_page(page: Any) -> str:
    """Rasterise one page and OCR it. Returns "" if OCR is unavailable."""
    try:
        import io

        import pytesseract  # noqa: F401 — presence check
        from PIL import Image

        pixmap = page.get_pixmap(dpi=200)
        image = Image.open(io.BytesIO(pixmap.tobytes("png")))
        return pytesseract.image_to_string(image) or ""
    except Exception as exc:  # noqa: BLE001 — OCR is best-effort by design
        _log.debug("vetting: OCR unavailable/failed for a page: %s", exc)
        return ""


def _extract_image(content: bytes) -> str:
    """A photographed document — the common case on a phone."""
    try:
        import io

        import pytesseract
        from PIL import Image

        return (pytesseract.image_to_string(Image.open(io.BytesIO(content)))
                or "")[:_MAX_EXTRACT_CHARS]
    except Exception as exc:  # noqa: BLE001
        _log.debug("vetting: image OCR failed: %s", exc)
        return ""


def _extract_docx(content: bytes) -> str:
    try:
        import io

        import docx

        document = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in document.paragraphs)[:_MAX_EXTRACT_CHARS]
    except Exception as exc:  # noqa: BLE001
        _log.debug("vetting: docx extraction failed: %s", exc)
        return ""


def _extract_email(content: bytes) -> str:
    """An .eml, INCLUDING its attachments.

    For a referee reply the reference is almost always the attachment and the
    covering note says "please find attached" — reading only the body would
    miss the evidence entirely, which is the whole reason this walks parts.
    """
    try:
        from email import policy
        from email.parser import BytesParser

        message = BytesParser(policy=policy.default).parsebytes(content)
        parts: list[str] = []
        for header in ("Subject", "From", "To", "Date"):
            value = message.get(header)
            if value:
                parts.append(f"{header}: {value}")

        for part in message.walk():
            if part.is_multipart():
                continue
            name = part.get_filename() or ""
            ctype = (part.get_content_type() or "").lower()
            if ctype in ("text/plain", "text/html") and not name:
                try:
                    parts.append(part.get_content())
                except Exception:  # noqa: BLE001
                    pass
                continue
            payload = part.get_payload(decode=True)
            if payload and name:
                # Recurse ONE level into the attachment by its own filename.
                nested = decode_text_best_effort(payload, name)
                if nested:
                    parts.append(f"[attachment: {name}]\n{nested}")
        return "\n".join(parts)[:_MAX_EXTRACT_CHARS]
    except Exception as exc:  # noqa: BLE001
        _log.debug("vetting: email extraction failed: %s", exc)
        return ""


def decode_text_best_effort(content: bytes, filename: str) -> str:
    """Extract text we can actually read, without pretending we read more.

    R-F3265 — this used to return "" for anything that was not .txt/.csv/.md/
    .json, which is to say for every document a real applicant sends. PDFs,
    scans, phone photos, DOCX references and emailed replies all reported
    `extraction_unavailable`, landed as OTHER, satisfied no requirement and
    evidenced no period. The dependencies to read them (PyMuPDF, tesseract,
    python-docx) were already in the image, and `intel/pdf_deep_ingest.py` had
    been using them for the DD side the whole time.

    What has NOT changed: this only ever supplies TEXT. Everything downstream
    still treats extraction as a PROPOSAL — the confidence floor, the
    authenticity flags and the RECEIVED-vs-ACCEPTED distinction are untouched,
    and an unreadable document still returns "" so it routes to a human as a
    DISCLOSED gap rather than a silent pass. Reading a date off a payslip is
    not verifying an engagement.
    """
    if not content:
        return ""
    lowered = filename.lower()
    try:
        if lowered.endswith(".json"):
            try:
                return json.dumps(json.loads(content))[:_MAX_EXTRACT_CHARS]
            except Exception:  # noqa: BLE001 — fall through to plain decode
                pass
        if lowered.endswith((".txt", ".csv", ".md", ".json", ".log")):
            return content.decode("utf-8", errors="replace")[:_MAX_EXTRACT_CHARS]
        if lowered.endswith(".pdf"):
            return _extract_pdf(content)
        if lowered.endswith((".docx", ".dotx")):
            return _extract_docx(content)
        if lowered.endswith((".eml", ".msg")):
            return _extract_email(content)
        if lowered.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff",
                             ".bmp", ".webp", ".heic")):
            return _extract_image(content)
    except Exception as exc:  # noqa: BLE001 — an upload must never fail on decode
        _log.debug("vetting: extraction failed for %s: %s", filename, exc)
    return ""

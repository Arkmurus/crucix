"""R-F945 + R-F946 — the live "she only reviewed the first clauses / reviews are
slow" root causes, found 2026-05-27 on the Korvera UTS Master Agency Agreement.

R-F945 (contention): the 7-layer retrieval context fed the FULL message —
including a 60K-char [ATTACHED DOCUMENT] body — to search_knowledge (a GIL-bound
scan of ~45K facts) and the embedder, freezing the event loop 5s+ between every
review step (wedge_674). Fix: retrieval layers search a stripped/capped QUERY;
the full document still reaches the LLM via the user_prompt.

R-F946 (extraction): _read_docx read ALL paragraphs then APPENDED all table
cells at the end, so a contract with interleaved tables came out jumbled — the
review stopped mid-sentence where a table interrupted and later clauses landed
out of order. Fix: walk the document body in reading order.
"""
from __future__ import annotations

import os
import tempfile

import pytest

from aria_service.aria_engine import _context_search_query
from aria_service.intel.document_reader import _read_docx


# ── R-F945 — retrieval query is stripped + capped ────────────────────────────

def test_rf945_strips_document_block_from_query():
    doc = "[ATTACHED DOCUMENT: c.docx]\n" + ("CLAUSE BODY " * 4000) + "\n[END ATTACHED DOCUMENT]"
    q = _context_search_query("Please review this agreement.\n" + doc)
    assert "CLAUSE BODY" not in q
    assert "review this agreement" in q.lower()
    assert len(q) <= 1500


def test_rf945_caps_long_plain_message():
    q = _context_search_query("x" * 50000)
    assert len(q) <= 1500


def test_rf945_short_message_unchanged():
    # normal chat must be unaffected — same string back
    assert _context_search_query("what are the biggest defence tenders in Jordan?") == \
        "what are the biggest defence tenders in Jordan?"


def test_rf945_empty_safe():
    assert _context_search_query("") == ""
    assert _context_search_query(None) in ("", None)


# ── R-F946 — .docx extracted in document order (table between paragraphs) ─────

def _make_docx_with_interleaved_table(path):
    from docx import Document
    d = Document()
    # Pad past the 100-char raw-parse floor so the in-order path is exercised
    # (a real contract is tens of K chars; a 3-word doc would hit the fallback).
    d.add_paragraph("ALPHA_FIRST_PARAGRAPH " + "lorem ipsum dolor sit amet " * 3)
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "BETA_TABLE_CELL " + "consectetur adipiscing elit " * 3
    d.add_paragraph("GAMMA_LAST_PARAGRAPH " + "sed do eiusmod tempor " * 3)
    d.save(path)


def test_rf946_docx_extracted_in_reading_order():
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        _make_docx_with_interleaved_table(path)
        r = _read_docx(path)
        t = r.text or ""
        assert "ALPHA_FIRST_PARAGRAPH" in t
        assert "BETA_TABLE_CELL" in t
        assert "GAMMA_LAST_PARAGRAPH" in t
        # THE FIX: the table cell appears BETWEEN the two paragraphs (reading
        # order), not appended after the last one (the old jumbled behaviour).
        assert t.index("ALPHA_FIRST_PARAGRAPH") < t.index("BETA_TABLE_CELL") < t.index("GAMMA_LAST_PARAGRAPH")
        assert r.confidence >= 0.9
    finally:
        os.remove(path)


def test_rf946_plain_paragraph_docx():
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        from docx import Document
        d = Document()
        for i in range(5):
            d.add_paragraph(f"PARAGRAPH_NUMBER_{i}")
        d.save(path)
        r = _read_docx(path)
        for i in range(5):
            assert f"PARAGRAPH_NUMBER_{i}" in r.text
    finally:
        os.remove(path)
